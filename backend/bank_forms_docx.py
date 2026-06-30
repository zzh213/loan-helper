"""一键生成可直接提交银行的成品材料 Word(.docx)版,内容可二次编辑:
  1) 授信申请表  2) 经营情况说明书  3) 财政贴息申报台账
带企业 Logo 占位、签字盖章栏,字段自动填充。
"""
import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from models import EnterpriseProfile, RecommendResponse

CREDIT_CN = {"excellent": "优秀", "good": "良好", "fair": "一般", "poor": "较差"}
PURPOSE_CN = {
    "working_capital": "流动资金周转", "equipment": "设备采购", "expansion": "扩大经营",
    "inventory": "备货采购", "rd": "研发投入", "other": "其他",
}
BANK_RGB = {
    "国有银行": (0x9a, 0x1f, 0x1f), "股份制银行": (0x1e, 0x40, 0xaf),
    "城商行": (0x0f, 0x76, 0x6e), "政策性/普惠": (0xb4, 0x53, 0x09),
    "互联网银行": (0x6d, 0x28, 0xb9), "小额贷款公司": (0x37, 0x41, 0x51),
}


def _title(doc, text, rgb):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("【企业Logo】  " + text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(*rgb)


def _sub(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)


def _h2(doc, text, rgb):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(*rgb)


def _kv(doc, rows):
    t = doc.add_table(rows=len(rows), cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, r in enumerate(rows):
        for j, v in enumerate(r):
            t.cell(i, j).text = str(v)
    return t


def _sign(doc, today):
    doc.add_paragraph()
    doc.add_paragraph("企业(公章):________________      法定代表人(签字):________________")
    doc.add_paragraph(f"申报日期:{today}            联系电话:________________")


def build_bank_package_docx(profile: EnterpriseProfile, result: RecommendResponse, hidden=None, bank_name="") -> bytes:
    best = result.plans[0] if result.plans else None
    rgb = BANK_RGB.get(best.provider_type if best else "", (0x1e, 0x40, 0xaf))
    company = profile.company_name or "(请填写企业名称)"
    today = datetime.now().strftime("%Y 年 %m 月 %d 日")
    doc = Document()

    # 1. 授信申请表
    if bank_name:
        _h2(doc, f"致:{bank_name} 普惠金融部", rgb)
    _title(doc, "企业授信申请表", rgb)
    _sub(doc, f"申请机构:{best.provider_type if best else '——'} · "
              f"拟申请产品:{best.product_name if best else '——'}")
    _h2(doc, "一、申请人基本信息", rgb)
    _kv(doc, [
        ["企业名称", profile.company_name or "", "所属行业", profile.industry],
        ["经营年限", f"{profile.years_in_business} 年", "员工人数", f"{profile.employees} 人"],
        ["注册资本", f"{profile.registered_capital} 万元", "年营业额", f"{profile.annual_revenue} 万元"],
        ["征信状况", CREDIT_CN.get(profile.credit_level.value, "-"),
         "纳税记录", "有" if profile.has_tax_record else "无"],
        ["开票流水", "有" if profile.has_invoice else "无",
         "当前逾期", "有" if profile.has_overdue else "无"],
    ])
    amount = best.estimated_amount if best else profile.loan_amount
    rate = f"{best.annual_rate_min}%-{best.annual_rate_max}%" if best else "——"
    term = best.suggested_term_months if best else profile.preferred_term_months
    monthly = best.monthly_payment_estimate if best else "——"
    _h2(doc, "二、授信申请要素", rgb)
    _kv(doc, [
        ["申请额度", f"{amount} 万元", "申请期限", f"{term} 个月"],
        ["资金用途", PURPOSE_CN.get(profile.loan_purpose.value, "-"), "预计年化利率", rate],
        ["还款方式", "等额本息", "预估月供", f"{monthly} 万元"],
        ["担保方式", ("抵押(估值 %s 万元)" % profile.collateral_value) if profile.has_collateral else "信用",
         "预计放款", best.expected_release_days if best else "——"],
    ])
    _h2(doc, "三、申请说明", rgb)
    doc.add_paragraph(
        f"本企业因{PURPOSE_CN.get(profile.loan_purpose.value, '经营周转')}需要,"
        f"特向贵行申请授信额度人民币 {amount} 万元,期限 {term} 个月。"
        "本企业经营状况良好,信用记录正常,具备相应还款能力,恳请贵行审批为盼。")
    _sign(doc, today)
    doc.add_page_break()

    # 2. 经营情况说明书
    _title(doc, "企业经营情况说明书", rgb)
    _sub(doc, company)
    _h2(doc, "一、企业基本情况", rgb)
    doc.add_paragraph(
        f"我单位 {company},属{profile.industry}行业,成立至今已稳定经营 "
        f"{profile.years_in_business} 年,注册资本 {profile.registered_capital} 万元,"
        f"现有员工 {profile.employees} 人。近一年实现营业额约 {profile.annual_revenue} 万元,"
        "经营稳健,现金流良好。")
    _h2(doc, "二、信用与财务状况", rgb)
    doc.add_paragraph(
        "企业纳税记录" + ("规范、连续" if profile.has_tax_record else "正在完善") + ","
        "开票流水" + ("稳定" if profile.has_invoice else "持续积累中") + ","
        "征信状况" + CREDIT_CN.get(profile.credit_level.value, "良好") + ","
        + ("当前无逾期记录。" if not profile.has_overdue else "正积极处理历史逾期。"))
    _h2(doc, "三、融资需求与还款安排", rgb)
    doc.add_paragraph(
        f"本次拟向{best.provider_type if best else '银行'}申请融资 {amount} 万元,"
        f"主要用于{PURPOSE_CN.get(profile.loan_purpose.value, '经营周转')},"
        "预计可有效缓解资金压力、扩大经营规模、提升盈利能力。还款来源为日常经营收入,"
        "还款能力充足,风险可控。")
    if result.profile_highlights:
        _h2(doc, "四、企业优势", rgb)
        for h in result.profile_highlights:
            doc.add_paragraph(h, style="List Bullet")
    _sign(doc, today)
    doc.add_page_break()

    # 3. 贴息申报台账
    _title(doc, "财政贴息 / 扶持政策申报台账", rgb)
    _sub(doc, f"{company} · {today}")
    if result.subsidies:
        t = doc.add_table(rows=1, cols=6)
        t.style = "Table Grid"
        hdr = ["序号", "政策名称", "主管部门", "扶持内容", "申报要点", "状态"]
        for j, v in enumerate(hdr):
            t.cell(0, j).text = v
        for i, s in enumerate(result.subsidies, 1):
            c = t.add_row().cells
            c[0].text = str(i); c[1].text = s.name; c[2].text = s.authority
            c[3].text = s.benefit; c[4].text = s.apply_points; c[5].text = "待申报"
        doc.add_paragraph(f"共匹配 {len(result.subsidies)} 项可申报政策,请按各主管部门要求准备材料并在窗口期内提交。")
    else:
        doc.add_paragraph("暂未匹配到可申报政策,建议完善纳税、参保等资质后再行申报。")

    if hidden:
        _h2(doc, "园区 / 区县独家隐藏贴息(非公开,需定向申报)", rgb)
        ht = doc.add_table(rows=1, cols=6)
        ht.style = "Table Grid"
        hh = ["序号", "政策名称", "所在地区", "最高额", "贴息", "扶持内容/申报路径"]
        for j, v in enumerate(hh):
            ht.cell(0, j).text = v
        for i, s in enumerate(hidden, 1):
            c = ht.add_row().cells
            c[0].text = str(i); c[1].text = s.get("name", ""); c[2].text = s.get("region", "")
            c[3].text = f"{s.get('amount_max','-')}万"; c[4].text = f"{s.get('rate_subsidy','-')}%"
            c[5].text = f"{s.get('benefit','')}\n申报:{s.get('apply_points','')}"
        doc.add_paragraph(f"共解锁 {len(hidden)} 项独家隐藏贴息,名额有限,建议尽快向对应园区/区县窗口定向申报。")
    _sign(doc, today)
    doc.add_paragraph("说明:本台账金额、政策为系统估算与匹配示意,最终以主管部门核定为准。")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
